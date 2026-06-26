from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "baocao" / "bao_cao_PhamPhiLong_2211074.docx"
OUT_DIR = ROOT / "baocao" / "chapter_1_diagrams"

FONT_REG = Path("C:/Windows/Fonts/arial.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/arialbd.ttf")

COLORS = {
    "bg": "#F8FAFC",
    "ink": "#111827",
    "muted": "#475569",
    "border": "#CBD5E1",
    "blue": "#2563EB",
    "blue_light": "#DBEAFE",
    "green": "#059669",
    "green_light": "#D1FAE5",
    "amber": "#D97706",
    "amber_light": "#FEF3C7",
    "violet": "#7C3AED",
    "violet_light": "#EDE9FE",
    "red": "#DC2626",
    "red_light": "#FEE2E2",
    "slate_light": "#E2E8F0",
    "white": "#FFFFFF",
}


def font(size, bold=False):
    return ImageFont.truetype(str(FONT_BOLD if bold else FONT_REG), size)


def text_size(draw, text, fnt):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw, text, fnt, max_width):
    words = str(text).split()
    lines = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if text_size(draw, candidate, fnt)[0] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(draw, text, xy, fnt, fill, max_width, line_gap=8, center=False):
    x, y = xy
    lines = wrap_text(draw, text, fnt, max_width)
    line_h = text_size(draw, "Ag", fnt)[1] + line_gap
    for line in lines:
        tx = x
        if center:
            tx = x + (max_width - text_size(draw, line, fnt)[0]) / 2
        draw.text((tx, y), line, font=fnt, fill=fill)
        y += line_h
    return y


def rounded(draw, box, fill, outline=None, width=3, radius=24):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#334155", width=5, label=None, label_pos=0.5, label_bg="#FFFFFF"):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=fill, width=width)
    angle = math.atan2(y2 - y1, x2 - x1)
    head_len = 22
    head_w = 12
    p1 = (
        x2 - head_len * math.cos(angle) + head_w * math.sin(angle),
        y2 - head_len * math.sin(angle) - head_w * math.cos(angle),
    )
    p2 = (
        x2 - head_len * math.cos(angle) - head_w * math.sin(angle),
        y2 - head_len * math.sin(angle) + head_w * math.cos(angle),
    )
    draw.polygon([(x2, y2), p1, p2], fill=fill)
    if label:
        fnt = font(24, True)
        lx = x1 + (x2 - x1) * label_pos
        ly = y1 + (y2 - y1) * label_pos
        tw, th = text_size(draw, label, fnt)
        pad = 10
        rounded(draw, (lx - tw / 2 - pad, ly - th / 2 - pad, lx + tw / 2 + pad, ly + th / 2 + pad), label_bg, fill, 2, 12)
        draw.text((lx - tw / 2, ly - th / 2 - 2), label, font=fnt, fill=fill)


def poly_arrow(draw, points, fill="#334155", width=5, label=None, label_xy=None, label_bg="#FFFFFF"):
    for p1, p2 in zip(points, points[1:]):
        draw.line((*p1, *p2), fill=fill, width=width)
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    angle = math.atan2(y2 - y1, x2 - x1)
    head_len = 22
    head_w = 12
    p1 = (
        x2 - head_len * math.cos(angle) + head_w * math.sin(angle),
        y2 - head_len * math.sin(angle) - head_w * math.cos(angle),
    )
    p2 = (
        x2 - head_len * math.cos(angle) - head_w * math.sin(angle),
        y2 - head_len * math.sin(angle) + head_w * math.cos(angle),
    )
    draw.polygon([(x2, y2), p1, p2], fill=fill)
    if label and label_xy:
        fnt = font(24, True)
        lx, ly = label_xy
        tw, th = text_size(draw, label, fnt)
        pad = 10
        rounded(draw, (lx - tw / 2 - pad, ly - th / 2 - pad, lx + tw / 2 + pad, ly + th / 2 + pad), label_bg, fill, 2, 12)
        draw.text((lx - tw / 2, ly - th / 2 - 2), label, font=fnt, fill=fill)


def title(draw, main, subtitle=None, w=2400):
    draw.text((90, 60), main, font=font(50, True), fill=COLORS["ink"])
    if subtitle:
        draw.text((92, 124), subtitle, font=font(27), fill=COLORS["muted"])
    draw.line((90, 175, w - 90, 175), fill=COLORS["border"], width=3)


def box_with_items(draw, box, heading, items, fill, accent, heading_size=30, item_size=24):
    x1, y1, x2, y2 = box
    rounded(draw, box, fill, accent, 4, 28)
    draw.text((x1 + 28, y1 + 24), heading, font=font(heading_size, True), fill=COLORS["ink"])
    y = y1 + 78
    for item in items:
        draw.ellipse((x1 + 34, y + 11, x1 + 46, y + 23), fill=accent)
        y = draw_wrapped(draw, item, (x1 + 60, y), font(item_size), COLORS["muted"], x2 - x1 - 92, 6)
        y += 8


def draw_client_server(path):
    w, h = 2400, 1450
    img = Image.new("RGB", (w, h), COLORS["bg"])
    draw = ImageDraw.Draw(img)
    title(draw, "Kiến trúc Client - Server của hệ thống LingoConnect", "Phân tách rõ Frontend, Backend API, dữ liệu và dịch vụ ngoài", w)

    lanes = [
        (80, 230, 650, 1290, "CLIENT", COLORS["blue_light"], COLORS["blue"]),
        (705, 230, 1545, 1290, "BACKEND API", COLORS["green_light"], COLORS["green"]),
        (1600, 230, 2320, 1290, "DATA & EXTERNAL SERVICES", COLORS["violet_light"], COLORS["violet"]),
    ]
    for x1, y1, x2, y2, name, fill, accent in lanes:
        rounded(draw, (x1, y1, x2, y2), fill, accent, 4, 30)
        draw.text((x1 + 28, y1 + 24), name, font=font(28, True), fill=accent)

    box_with_items(
        draw,
        (130, 330, 600, 575),
        "Người học / Admin",
        ["Truy cập bằng trình duyệt", "Thao tác học tập, làm bài, quản trị nội dung", "Nhận phản hồi, EXP, trạng thái mở khóa"],
        COLORS["white"],
        COLORS["blue"],
    )
    box_with_items(
        draw,
        (130, 675, 600, 1110),
        "React SPA",
        ["React Router điều hướng trang", "AuthContext giữ phiên đăng nhập", "Axios client gắn JWT và chuẩn hóa lỗi", "UI cho Listening, Reading, Speaking, Writing, Grammar, Vocabulary, Admin"],
        COLORS["white"],
        COLORS["blue"],
    )

    box_with_items(
        draw,
        (755, 320, 1495, 520),
        "Express API Layer",
        ["Route prefix /api/v1/*", "JWT middleware, role guard, validation", "Controller nhận request và trả JSON response"],
        COLORS["white"],
        COLORS["green"],
        32,
        25,
    )
    modules = [
        ("Auth / User", 760, 610, COLORS["blue_light"], COLORS["blue"]),
        ("Learning Skills", 1015, 610, COLORS["green_light"], COLORS["green"]),
        ("Vocabulary", 1270, 610, COLORS["amber_light"], COLORS["amber"]),
        ("Games & Tasks", 760, 785, COLORS["violet_light"], COLORS["violet"]),
        ("Billing Plus", 1015, 785, COLORS["red_light"], COLORS["red"]),
        ("Admin CRUD", 1270, 785, COLORS["slate_light"], "#334155"),
    ]
    for label, x, y, fill, accent in modules:
        rounded(draw, (x, y, x + 210, y + 115), fill, accent, 3, 22)
        draw_wrapped(draw, label, (x + 18, y + 32), font(25, True), COLORS["ink"], 174, 4, True)
    box_with_items(
        draw,
        (815, 1010, 1435, 1210),
        "Service / Repository",
        ["Xử lý nghiệp vụ, chấm điểm, mở khóa bài tiếp theo", "Truy vấn dữ liệu và gọi dịch vụ AI/thanh toán khi cần"],
        COLORS["white"],
        COLORS["green"],
        30,
        24,
    )

    box_with_items(
        draw,
        (1660, 315, 2260, 590),
        "PostgreSQL",
        ["Users, Roles, Subscription", "Lessons, Questions, Progress", "Collections, MiniGame, DailyTasks, UserStats"],
        COLORS["white"],
        COLORS["violet"],
        32,
        25,
    )
    box_with_items(
        draw,
        (1660, 665, 1945, 880),
        "Flask ASR",
        ["faster-whisper", "Nhận audio ghi âm", "Trả transcript"],
        COLORS["white"],
        COLORS["amber"],
        28,
        23,
    )
    box_with_items(
        draw,
        (1975, 665, 2260, 880),
        "LLM",
        ["Sinh bài Speaking AI", "Chấm Writing linh hoạt", "Cache / timeout"],
        COLORS["white"],
        COLORS["red"],
        28,
        23,
    )
    box_with_items(
        draw,
        (1660, 970, 2260, 1195),
        "Nền tảng triển khai",
        ["Frontend: Vercel", "Backend & database: Railway", "Webhook thanh toán: SePay callback"],
        COLORS["white"],
        "#334155",
        30,
        24,
    )

    arrow(draw, (365, 575), (365, 675), COLORS["blue"], 5, "mở SPA")
    poly_arrow(draw, [(600, 850), (675, 850), (675, 420), (755, 420)], COLORS["blue"], 6, "REST + JWT", (675, 640))
    arrow(draw, (1495, 420), (1660, 455), COLORS["violet"], 6, "SQL")
    arrow(draw, (1435, 1085), (1660, 455), COLORS["violet"], 6)
    arrow(draw, (1435, 1060), (1660, 770), COLORS["amber"], 6, "audio", 0.72)
    poly_arrow(draw, [(1435, 1135), (1565, 1135), (1565, 630), (1975, 630), (1975, 665)], COLORS["red"], 6, "prompt", (1565, 900))
    poly_arrow(draw, [(1660, 1080), (1545, 1080), (1545, 1160), (1435, 1160)], "#334155", 5, "webhook", (1545, 1118))
    poly_arrow(draw, [(755, 485), (650, 485), (650, 720), (600, 720)], COLORS["green"], 5, "JSON response", (650, 605))

    draw.text((90, 1340), "Luồng chính: Browser gửi request có JWT → Express kiểm tra quyền → module nghiệp vụ xử lý → PostgreSQL/dịch vụ ngoài → trả JSON để React cập nhật UI.", font=font(26, True), fill=COLORS["ink"])
    img.save(path, quality=95)


def draw_spa_rest(path):
    w, h = 2400, 1450
    img = Image.new("RGB", (w, h), COLORS["bg"])
    draw = ImageDraw.Draw(img)
    title(draw, "Mô hình SPA kết hợp RESTful API", "Điều hướng phía client, giao tiếp dữ liệu qua tài nguyên REST", w)

    swimlanes = [
        (80, 240, 2320, 565, "FRONTEND - React Single Page Application", COLORS["blue_light"], COLORS["blue"]),
        (80, 610, 2320, 995, "BACKEND - Node.js/Express REST API", COLORS["green_light"], COLORS["green"]),
        (80, 1040, 2320, 1300, "PERSISTENCE / EXTERNAL", COLORS["violet_light"], COLORS["violet"]),
    ]
    for x1, y1, x2, y2, label, fill, accent in swimlanes:
        rounded(draw, (x1, y1, x2, y2), fill, accent, 4, 24)
        draw.text((x1 + 25, y1 + 18), label, font=font(26, True), fill=accent)

    steps = [
        ("1", "URL / thao tác người dùng", 140, 345, COLORS["blue"]),
        ("2", "React Router chọn page", 480, 345, COLORS["blue"]),
        ("3", "Component đọc state", 820, 345, COLORS["blue"]),
        ("4", "AuthContext kiểm tra JWT", 1160, 345, COLORS["blue"]),
        ("5", "Axios client gửi API", 1500, 345, COLORS["blue"]),
        ("6", "UI cập nhật từ JSON", 1840, 345, COLORS["blue"]),
    ]
    for num, label, x, y, accent in steps:
        rounded(draw, (x, y, x + 260, y + 120), COLORS["white"], accent, 3, 20)
        draw.ellipse((x + 18, y + 22, x + 62, y + 66), fill=accent)
        draw.text((x + 32, y + 27), num, font=font(24, True), fill=COLORS["white"])
        draw_wrapped(draw, label, (x + 78, y + 24), font(24, True), COLORS["ink"], 160, 5)
    for (_, _, x, y, _), (_, _, x2, y2, _) in zip(steps, steps[1:]):
        arrow(draw, (x + 260, y + 60), (x2, y2 + 60), COLORS["blue"], 4)

    resources = [
        ("GET", "/api/v1/auth/me", "lấy phiên đăng nhập"),
        ("POST", "/api/v1/speaking/progress", "lưu tiến độ nói"),
        ("GET", "/api/v1/listening/lessons", "danh sách bài học"),
        ("POST", "/api/v1/collections/submissions", "gửi học phần công khai"),
        ("PUT", "/api/v1/admin/*", "cập nhật nội dung"),
    ]
    rx, ry = 150, 715
    draw.text((rx, ry - 28), "REST resources", font=font(29, True), fill=COLORS["green"])
    for i, (method, route, desc) in enumerate(resources):
        y = ry + i * 52
        rounded(draw, (rx, y, rx + 104, y + 36), COLORS["white"], COLORS["green"], 2, 12)
        draw.text((rx + 18, y + 7), method, font=font(19, True), fill=COLORS["green"])
        draw.text((rx + 125, y + 5), route, font=font(23, True), fill=COLORS["ink"])
        draw.text((rx + 600, y + 6), desc, font=font(22), fill=COLORS["muted"])

    backend_steps = [
        ("Route", "Định tuyến /api/v1 theo module", 1020, 690, COLORS["green"]),
        ("Middleware", "JWT, role, validation, upload", 1325, 690, COLORS["amber"]),
        ("Controller", "Nhận request, chuẩn hóa response", 1630, 690, COLORS["violet"]),
        ("Service", "Nghiệp vụ, chấm điểm, mở khóa", 1935, 690, COLORS["red"]),
    ]
    for title_text, body, x, y, accent in backend_steps:
        rounded(draw, (x, y, x + 250, y + 150), COLORS["white"], accent, 3, 20)
        draw.text((x + 22, y + 22), title_text, font=font(26, True), fill=accent)
        draw_wrapped(draw, body, (x + 22, y + 64), font(22), COLORS["muted"], 206, 5)
    for left, right in zip(backend_steps, backend_steps[1:]):
        arrow(draw, (left[2] + 250, left[3] + 75), (right[2], right[3] + 75), COLORS["green"], 4)

    data_boxes = [
        ("PostgreSQL", "Trạng thái bền vững: users, lessons, progress, billing", 220, 1115, COLORS["violet"]),
        ("Flask / Whisper", "Multipart audio → transcript cho Speaking", 880, 1115, COLORS["amber"]),
        ("LLM API", "Sinh bài cá nhân hóa, chấm Writing có kiểm soát", 1420, 1115, COLORS["red"]),
        ("SePay", "Webhook thanh toán Plus", 1930, 1115, "#334155"),
    ]
    for head, body, x, y, accent in data_boxes:
        rounded(draw, (x, y, x + 410, y + 120), COLORS["white"], accent, 3, 20)
        draw.text((x + 22, y + 18), head, font=font(25, True), fill=accent)
        draw_wrapped(draw, body, (x + 22, y + 58), font(21), COLORS["muted"], 360, 4)

    arrow(draw, (1630, 465), (1630, 690), COLORS["blue"], 5, "HTTP request")
    arrow(draw, (2060, 840), (2060, 1115), COLORS["green"], 5, "query/call")
    arrow(draw, (1970, 1115), (1970, 840), COLORS["violet"], 5, "result")
    arrow(draw, (1840, 760), (1840, 465), COLORS["green"], 5, "JSON")

    draw.text((95, 1350), "Đặc điểm chính: frontend không tải lại toàn trang; backend expose API tài nguyên; mọi thao tác học tập/quản trị đi qua JWT và trả response JSON nhất quán.", font=font(26, True), fill=COLORS["ink"])
    img.save(path, quality=95)


def draw_fuzzy_alignment(path):
    w, h = 2400, 1650
    img = Image.new("RGB", (w, h), COLORS["bg"])
    draw = ImageDraw.Draw(img)
    title(draw, "Sơ đồ khối thuật toán Fuzzy Word Alignment", "So sánh transcript ASR với câu mẫu để chấm phát âm linh hoạt", w)

    pipeline = [
        ("1", "Audio ghi âm", "File WAV/WebM từ trình duyệt", 100, 280, COLORS["blue"]),
        ("2", "ASR transcript", "faster-whisper chuyển giọng nói thành văn bản", 455, 280, COLORS["amber"]),
        ("3", "Chuẩn hóa văn bản", "lowercase, bỏ dấu câu, mở rộng contractions, số → chữ, bỏ filler words", 810, 280, COLORS["green"]),
        ("4", "Token hóa", "Tách câu mẫu và transcript thành chuỗi từ", 1165, 280, COLORS["violet"]),
        ("5", "Ma trận tương đồng", "Levenshtein similarity giữa từng cặp từ", 1520, 280, COLORS["red"]),
        ("6", "Căn chỉnh động", "Giữ thứ tự từ, cho phép thiếu/thừa/thay thế", 1875, 280, "#334155"),
    ]
    for num, head, body, x, y, accent in pipeline:
        rounded(draw, (x, y, x + 300, y + 185), COLORS["white"], accent, 4, 24)
        draw.ellipse((x + 20, y + 22, x + 68, y + 70), fill=accent)
        draw.text((x + 36, y + 29), num, font=font(24, True), fill=COLORS["white"])
        draw.text((x + 84, y + 25), head, font=font(25, True), fill=accent)
        draw_wrapped(draw, body, (x + 24, y + 86), font(21), COLORS["muted"], 252, 5)
    for left, right in zip(pipeline, pipeline[1:]):
        arrow(draw, (left[3] + 300, left[4] + 92), (right[3], right[4] + 92), COLORS["ink"], 4)

    rounded(draw, (105, 565, 1135, 1160), COLORS["white"], COLORS["blue"], 4, 28)
    draw.text((140, 600), "Ví dụ dữ liệu đầu vào", font=font(32, True), fill=COLORS["blue"])
    draw.text((145, 680), "Câu mẫu:", font=font(26, True), fill=COLORS["ink"])
    draw.text((315, 680), "I would like to book a table for two.", font=font(27), fill=COLORS["muted"])
    draw.text((145, 745), "Transcript:", font=font(26, True), fill=COLORS["ink"])
    draw.text((315, 745), "I would like book table for to.", font=font(27), fill=COLORS["muted"])

    rows = [
        ("I", "I", "đúng", COLORS["green_light"], COLORS["green"]),
        ("would", "would", "đúng", COLORS["green_light"], COLORS["green"]),
        ("like", "like", "đúng", COLORS["green_light"], COLORS["green"]),
        ("to", "—", "thiếu", COLORS["red_light"], COLORS["red"]),
        ("book", "book", "đúng", COLORS["green_light"], COLORS["green"]),
        ("a", "—", "thiếu", COLORS["red_light"], COLORS["red"]),
        ("table", "table", "đúng", COLORS["green_light"], COLORS["green"]),
        ("for", "for", "đúng", COLORS["green_light"], COLORS["green"]),
        ("two", "to", "gần đúng/sai", COLORS["amber_light"], COLORS["amber"]),
    ]
    x0, y0 = 150, 835
    headers = ["Mẫu", "Người học", "Kết luận"]
    widths = [200, 250, 250]
    cx = x0
    for head, width in zip(headers, widths):
        rounded(draw, (cx, y0, cx + width, y0 + 48), COLORS["slate_light"], "#64748B", 2, 10)
        draw.text((cx + 20, y0 + 10), head, font=font(21, True), fill=COLORS["ink"])
        cx += width + 16
    y = y0 + 62
    for ref, hyp, status, fill, accent in rows[:4]:
        cx = x0
        for value, width in zip([ref, hyp, status], widths):
            rounded(draw, (cx, y, cx + width, y + 42), fill, accent, 2, 10)
            draw.text((cx + 20, y + 8), value, font=font(20, True), fill=accent)
            cx += width + 16
        y += 52
    draw.text((150, 1120), "Bảng ví dụ rút gọn; thuật toán thực tế xử lý toàn bộ chuỗi từ và giữ thứ tự xuất hiện.", font=font(22), fill=COLORS["muted"])

    rounded(draw, (1240, 565, 2295, 1010), COLORS["white"], COLORS["green"], 4, 28)
    draw.text((1275, 600), "Quy tắc chấm điểm", font=font(32, True), fill=COLORS["green"])
    scoring = [
        ("Matched", "Từ khớp chính xác hoặc similarity vượt ngưỡng."),
        ("Missing", "Từ trong câu mẫu không xuất hiện trong transcript."),
        ("Extra", "Từ người học nói thêm, không có trong câu mẫu."),
        ("Substituted", "Có từ tương ứng theo thứ tự nhưng độ giống thấp."),
        ("Score", "Điểm = matched weighted / số từ mục tiêu, có trừ lỗi thiếu/thừa."),
    ]
    y = 670
    for head, body in scoring:
        rounded(draw, (1280, y, 1510, y + 48), COLORS["green_light"], COLORS["green"], 2, 12)
        draw.text((1300, y + 11), head, font=font(21, True), fill=COLORS["green"])
        draw_wrapped(draw, body, (1535, y + 7), font(22), COLORS["muted"], 700, 4)
        y += 64

    rounded(draw, (220, 1190, 2180, 1490), COLORS["white"], COLORS["violet"], 4, 28)
    draw.text((255, 1222), "Kết quả trả về giao diện Speaking", font=font(32, True), fill=COLORS["violet"])
    outputs = [
        ("Điểm phát âm", "Hiển thị % đạt theo từng câu và toàn bài."),
        ("Từ đúng", "Highlight xanh để người học biết phần đã nói ổn."),
        ("Từ thiếu/sai", "Highlight đỏ/vàng kèm gợi ý luyện lại."),
        ("Mở khóa tiến độ", "Nếu vượt ngưỡng, lưu progress và mở bài kế tiếp."),
    ]
    for i, (head, body) in enumerate(outputs):
        x = 270 + i * 470
        rounded(draw, (x, 1300, x + 400, 1425), COLORS["violet_light"], COLORS["violet"], 3, 22)
        draw.text((x + 26, 1322), head, font=font(24, True), fill=COLORS["violet"])
        draw_wrapped(draw, body, (x + 26, 1362), font(21), COLORS["muted"], 340, 4)

    arrow(draw, (2030, 465), (1770, 565), COLORS["ink"], 4)
    arrow(draw, (1770, 1010), (1730, 1190), COLORS["violet"], 5, "feedback")
    draw.text((95, 1572), "Ý nghĩa: thuật toán không yêu cầu transcript giống tuyệt đối, phù hợp với lỗi phát âm tự nhiên của người học nhưng vẫn kiểm soát thứ tự và mức độ sai lệch.", font=font(26, True), fill=COLORS["ink"])
    img.save(path, quality=95)


def insert_paragraph_before(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def remove_empty_paragraph_after(paragraph):
    current = paragraph._p.getnext()
    while current is not None:
        p = Paragraph(current, paragraph._parent)
        if p.text.strip() or current.xpath(".//w:drawing") or current.xpath(".//pic:pic"):
            break
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt


def insert_image_before_caption(doc, caption_text, image_path):
    target = None
    for paragraph in doc.paragraphs:
        if caption_text in paragraph.text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError(f"Caption not found: {caption_text}")

    previous = target._p.getprevious()
    while previous is not None:
        prev_para = Paragraph(previous, target._parent)
        has_image = bool(previous.xpath(".//w:drawing") or previous.xpath(".//pic:pic"))
        if prev_para.text.strip() or not has_image:
            break
        older = previous.getprevious()
        previous.getparent().remove(previous)
        previous = older

    pic_para = insert_paragraph_before(target)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(6.7))
    remove_empty_paragraph_after(target)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    img1 = OUT_DIR / "1_3_client_server_architecture.png"
    img2 = OUT_DIR / "1_3_spa_rest_api.png"
    img3 = OUT_DIR / "1_4_fuzzy_word_alignment.png"

    draw_client_server(img1)
    draw_spa_rest(img2)
    draw_fuzzy_alignment(img3)

    doc = Document(REPORT)
    insert_image_before_caption(doc, "Kiến trúc Client-Server của hệ thống LingoConnect", img1)
    insert_image_before_caption(doc, "Mô hình SPA kết hợp RESTful API", img2)
    insert_image_before_caption(doc, "Sơ đồ khối thuật toán Fuzzy Word Alignment", img3)
    doc.save(REPORT)
    print(f"Generated and inserted diagrams into {REPORT}")
    print(img1)
    print(img2)
    print(img3)


if __name__ == "__main__":
    main()
